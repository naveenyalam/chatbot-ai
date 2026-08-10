import os
from typing import Dict, Any
import docx
from app.services.document.base import BaseExtractor

class DocxExtractor(BaseExtractor):
    def extract(self, file_path: str) -> Dict[str, Any]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        doc = docx.Document(file_path)
        parts = []
        
        # Parse paragraphs and headings
        for para in doc.paragraphs:
            val = para.text.strip()
            if val:
                parts.append(val)
                
        # Parse tables
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                # Filter empty rows
                cells_filtered = [c for c in cells if c]
                if cells_filtered:
                    parts.append(" | ".join(cells_filtered))
                    
        full_text = "\n\n".join(parts).strip()
        if not full_text:
            full_text = "This document is empty."
            
        return {
            "text": full_text,
            "pages": [{"page_number": 1, "text": full_text}],
            "metadata": {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
        }
